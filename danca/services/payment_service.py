import requests
import json
from django.conf import settings
from decimal import Decimal


class PaymentService:
    """Serviço para integração com gateways de pagamento"""
    
    @staticmethod
    def pagar_inscricao(inscricao, tipo_pagamento, dados_pagamento, gateway='pagseguro'):
        """
        Processa pagamento de inscrição
        
        Args:
            inscricao: Objeto Inscricao
            tipo_pagamento: 'credit_card' ou 'pix'
            dados_pagamento: dict com dados do pagamento
                Para cartão: {'numero', 'nome', 'cvv', 'mes', 'ano', 'parcelas'}
                Para PIX: {}
            gateway: 'pagseguro' ou 'infinitepay'
        
        Returns:
            dict com resultado do pagamento
        """
        if tipo_pagamento == 'credit_card':
            if gateway == 'pagseguro':
                return PaymentService._pagseguro_card_payment(inscricao, dados_pagamento)
            elif gateway == 'infinitepay':
                return PaymentService._infinitepay_card_payment(inscricao, dados_pagamento)
            else:
                raise ValueError(f"Gateway não suportado: {gateway}")
        elif tipo_pagamento == 'pix':
            if gateway == 'pagseguro':
                return PaymentService._pagseguro_pix_payment(inscricao, dados_pagamento)
            elif gateway == 'infinitepay':
                return PaymentService._infinitepay_pix_payment(inscricao, dados_pagamento)
            else:
                raise ValueError(f"Gateway não suportado: {gateway}")
        else:
            raise ValueError(f"Tipo de pagamento não suportado: {tipo_pagamento}")

    @staticmethod
    def pagar_pedido(pedido, tipo_pagamento, dados_pagamento, gateway='pagseguro'):
        """Processa pagamento de PedidoCamisa"""
        if tipo_pagamento == 'credit_card':
            # Simulação: aprovar diretamente e registrar pagamento
            from danca.models import Pagamento
            pagamento = Pagamento.objects.create(
                tipo_modelo='pedido',
                valor_pago=pedido.valor_venda,
                numero_parcela=dados_pagamento.get('parcelas', 1),
                status_pagamento='pago',
                gateway_pagamento=gateway,
                payment_method='credit_card',
                nome_cartao=dados_pagamento.get('nome'),
                ultimos_digitos=str(dados_pagamento.get('numero', ''))[-4:],
                bandeira_cartao='',
            )
            # ligar generic FK
            from django.contrib.contenttypes.models import ContentType
            pagamento.content_type = ContentType.objects.get_for_model(pedido.__class__)
            pagamento.object_id = pedido.id
            pagamento.save()

            # Atualizar status do pedido para "pago"
            pedido.status = 'pago'
            pedido.save(update_fields=['status'])
            
            # Diminuir quantidade de camisas disponíveis
            if pedido.camisa and pedido.camisa.quantidade > 0:
                pedido.camisa.quantidade -= 1
                pedido.camisa.save(update_fields=['quantidade'])
            
            # Criar entrada automática no caixa
            PaymentService._criar_entrada_automatica(pedido, pagamento)

            # Gerar nota fiscal simples
            nota_path = PaymentService._gerar_nota_fiscal('pedido', pedido, pagamento)
            return {
                'success': True,
                'pagamento_id': pagamento.id,
                'status': 'pago',
                'message': 'Pagamento aprovado',
                'nota_file': nota_path,
            }
        elif tipo_pagamento == 'pix':
            # Gera QRCode PIX parecido com inscrição, com valor do pedido
            from danca.models import Pagamento
            from datetime import datetime
            import qrcode, io, base64
            from django.conf import settings

            pagamento = Pagamento.objects.create(
                tipo_modelo='pedido',
                valor_pago=pedido.valor_venda,
                numero_parcela=1,
                status_pagamento='pendente',
                gateway_pagamento='pix_bancario',
                payment_method='pix',
            )
            from django.contrib.contenttypes.models import ContentType
            pagamento.content_type = ContentType.objects.get_for_model(pedido.__class__)
            pagamento.object_id = pedido.id
            pagamento.transaction_id = f"PIX_PED_{pedido.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
            pagamento.save()

            chave = getattr(settings, 'PIX_CHAVE', 'pix@danca.com')
            razao = getattr(settings, 'PIX_RAZAO_SOCIAL', 'Sistema Dança')[:25]
            cidade = getattr(settings, 'PIX_CIDADE', 'Joao Pessoa')[:25]
            emv = f"00020126580014BR.GOV.BCB.PIX0136{chave}520400005303986540{float(pedido.valor_venda):.2f}5802BR5909{razao:<25}6009{cidade:<25}62170510***{pagamento.transaction_id}6304"
            from zlib import crc32
            crc = format(crc32(emv.encode()), 'x').upper()
            emv += crc
            qr = qrcode.QRCode(version=1, box_size=10, border=5)
            qr.add_data(emv)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            qr_b64 = base64.b64encode(buffer.getvalue()).decode()
            return {
                'success': True,
                'pagamento_id': pagamento.id,
                'transaction_id': pagamento.transaction_id,
                'status': 'pendente',
                'payment_type': 'pix',
                'qr_code': emv,
                'qr_code_base64': qr_b64,
                'valor': str(pedido.valor_venda),
            }
        else:
            raise ValueError('Tipo de pagamento inválido')

    @staticmethod
    def _criar_entrada_automatica(obj, pagamento):
        """Cria entrada automática no caixa quando pagamento é aprovado"""
        from danca.models import Entrada
        from django.utils import timezone
        from django.contrib.contenttypes.models import ContentType
        
        # Criar entrada apenas se o pagamento foi aprovado
        if pagamento.status_pagamento == 'pago':
            try:
                if pagamento.tipo_modelo == 'pedido':
                    if hasattr(obj, 'camisa') and obj.camisa:
                        descricao = f"Venda de Camisa - Pedido #{obj.id} - {obj.camisa.descricao}"
                    else:
                        descricao = f"Venda de Camisa - Pedido #{obj.id}"
                elif pagamento.tipo_modelo == 'inscricao':
                    evento = getattr(obj, 'evento', None)
                    if evento:
                        evento_desc = getattr(evento, 'descricao', f"Evento #{getattr(evento, 'id', '')}")
                    else:
                        evento_desc = f"Inscrição #{obj.id}"
                    nome_cliente = getattr(obj, 'nome', 'Cliente')
                    descricao = f"Inscrição em Evento - {evento_desc} - {nome_cliente}"
                else:
                    descricao = f"Pagamento {pagamento.tipo_modelo} #{obj.id if hasattr(obj, 'id') else ''}"
                
                # Obter data do pagamento
                data_pagamento = pagamento.data_pagamento
                if data_pagamento is None:
                    data_pagamento = timezone.now()
                elif hasattr(data_pagamento, 'date'):
                    data_pagamento = data_pagamento.date()
                elif isinstance(data_pagamento, str):
                    from datetime import datetime
                    try:
                        data_pagamento = datetime.strptime(data_pagamento.split('T')[0], '%Y-%m-%d').date()
                    except:
                        data_pagamento = timezone.now().date()
                else:
                    data_pagamento = timezone.now().date()
                
                # Verificar se já existe entrada para este pagamento
                # Usar transaction_id ou ID do pagamento para evitar duplicatas
                descricao_unica = f"{descricao} (Pagamento #{pagamento.id})"
                entradas_existentes = Entrada.objects.filter(
                    descricao__contains=f"Pagamento #{pagamento.id}"
                )
                
                if not entradas_existentes.exists():
                    Entrada.objects.create(
                        descricao=descricao_unica,
                        valor=pagamento.valor_pago,
                        data=data_pagamento
                    )
            except Exception as e:
                # Log do erro mas não interromper o fluxo
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Erro ao criar entrada automática: {str(e)}")
                print(f"Erro ao criar entrada automática: {str(e)}")
    
    @staticmethod
    def _gerar_nota_fiscal(tipo, obj, pagamento):
        from docx import Document
        from django.conf import settings
        import os
        os.makedirs(os.path.join(settings.MEDIA_ROOT, 'notas'), exist_ok=True)
        file_path = os.path.join(settings.MEDIA_ROOT, 'notas', f'nota_{pagamento.id}.docx')
        doc = Document()
        doc.add_heading('Nota Fiscal - Sistema Dança', 0)
        doc.add_paragraph(f'Pagamento ID: {pagamento.id}')
        doc.add_paragraph(f'Tipo: {tipo}')
        if tipo == 'pedido':
            doc.add_paragraph(f'Pedido: {obj.id} - Camisa: {obj.camisa.descricao}')
            doc.add_paragraph(f'Valor: R$ {obj.valor_venda}')
        else:
            doc.add_paragraph(f'Valor: R$ {pagamento.valor_pago}')
        doc.save(file_path)
        return file_path
    
    @staticmethod
    def pagar_inscricao_por_cartao(inscricao, dados_cartao, gateway='pagseguro'):
        """Método de compatibilidade - mantém API antiga"""
        return PaymentService.pagar_inscricao(inscricao, 'credit_card', dados_cartao, gateway)
    
    @staticmethod
    def _pagseguro_card_payment(inscricao, dados_cartao):
        """Integração com PagSeguro"""
        from danca.models import Pagamento
        
        url = "https://sandbox.api.pagseguro.com/orders"  # Sandbox URL
        
        # Configurar autenticação (email + token sandbox)
        email = getattr(settings, 'PAGSEGURO_EMAIL', '')
        token = getattr(settings, 'PAGSEGURO_TOKEN', '')
        
        headers = {
            'Authorization': f'Bearer {token}',
            'Content-Type': 'application/json',
        }
        
        payload = {
            "reference_id": f"inscricao_{inscricao.id}",
            "customer": {
                "name": inscricao.nome,
                "email": f"inscrito{inscricao.id}@exemplo.com",  # Em produção usar email real
                "tax_id": inscricao.cpf if inscricao.cpf else None,
            },
            "items": [
                {
                    "reference_id": f"item_{inscricao.id}",
                    "name": f"Inscrição - {inscricao.nome}",
                    "quantity": 1,
                    "unit_amount": int(inscricao.valor_total * 100)  # Valor em centavos
                }
            ],
            "charges": [
                {
                    "reference_id": f"charge_{inscricao.id}",
                    "description": f"Pagamento inscrição ID {inscricao.id}",
                    "amount": {
                        "value": int(inscricao.valor_total * 100),
                        "currency": "BRL"
                    },
                    "payment_method": {
                        "type": "CREDIT_CARD",
                        "installments": dados_cartao.get('parcelas', 1),
                        "capture": True,
                        "card": {
                            "encrypted": dados_cartao.get('numero'),  # Deve ser criptografado via JS
                            "exp_month": dados_cartao.get('mes'),
                            "exp_year": dados_cartao.get('ano'),
                            "security_code": dados_cartao.get('cvv'),
                            "holder": {
                                "name": dados_cartao.get('nome')
                            }
                        }
                    }
                }
            ]
        }
        
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            # Criar registro de pagamento
            pagamento = Pagamento.objects.create(
                tipo_modelo='inscricao',
                content_type_id=8,  # ContentType de Inscricao
                object_id=inscricao.id,
                valor_pago=inscricao.valor_total,
                numero_parcela=dados_cartao.get('parcelas', 1),
                status_pagamento='pago' if data.get('status') == 'PAID' else 'pendente',
                gateway_pagamento='pagseguro',
                transaction_id=data.get('id'),
                payment_method='credit_card',
                nome_cartao=dados_cartao.get('nome'),
                ultimos_digitos=dados_cartao.get('numero', '')[-4:],
            )
            
            # Criar entrada automática se pagamento foi aprovado
            if pagamento.status_pagamento == 'pago':
                PaymentService._criar_entrada_automatica(inscricao, pagamento)
            
            return {
                'success': True,
                'pagamento_id': pagamento.id,
                'transaction_id': data.get('id'),
                'status': data.get('status'),
                'message': 'Pagamento processado com sucesso'
            }
            
        except requests.exceptions.RequestException as e:
            return {
                'success': False,
                'message': f'Erro ao processar pagamento: {str(e)}'
            }
    
    @staticmethod
    def _infinitepay_card_payment(inscricao, dados_cartao):
        """Integração com InfinitePay"""
        from danca.models import Pagamento
        
        url = "https://api.infinitepay.io.br/v1/orders"
        
        # Configurar autenticação
        api_key = getattr(settings, 'INFINITEPAY_API_KEY', '')
        
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json',
        }
        
        payload = {
            "order_reference_id": f"inscricao_{inscricao.id}",
            "customer": {
                "name": inscricao.nome,
                "email": f"inscrito{inscricao.id}@exemplo.com",
                "document": inscricao.cpf if inscricao.cpf else None,
            },
            "amount": float(inscricao.valor_total),
            "currency": "BRL",
            "payment_method": {
                "type": "credit_card",
                "card": {
                    "number": dados_cartao.get('numero'),
                    "cvv": dados_cartao.get('cvv'),
                    "expiry_month": dados_cartao.get('mes'),
                    "expiry_year": dados_cartao.get('ano'),
                    "holder_name": dados_cartao.get('nome')
                },
                "installments": dados_cartao.get('parcelas', 1)
            }
        }
        
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            data = response.json()
            
            # Criar registro de pagamento
            pagamento = Pagamento.objects.create(
                tipo_modelo='inscricao',
                content_type_id=8,
                object_id=inscricao.id,
                valor_pago=inscricao.valor_total,
                numero_parcela=dados_cartao.get('parcelas', 1),
                status_pagamento='pago' if data.get('status') == 'approved' else 'pendente',
                gateway_pagamento='infinitepay',
                transaction_id=data.get('id'),
                payment_method='credit_card',
                nome_cartao=dados_cartao.get('nome'),
                ultimos_digitos=dados_cartao.get('numero', '')[-4:],
            )
            
            # Criar entrada automática se pagamento foi aprovado
            if pagamento.status_pagamento == 'pago':
                PaymentService._criar_entrada_automatica(inscricao, pagamento)
            
            return {
                'success': True,
                'pagamento_id': pagamento.id,
                'transaction_id': data.get('id'),
                'status': data.get('status'),
                'message': 'Pagamento processado com sucesso'
            }
            
        except requests.exceptions.RequestException as e:
            return {
                'success': False,
                'message': f'Erro ao processar pagamento: {str(e)}'
            }
    
    @staticmethod
    def _pagseguro_pix_payment(inscricao, dados_pix=None):
        """Integração com PagSeguro - PIX"""
        from danca.models import Pagamento
        from datetime import datetime
        import qrcode
        import io
        import base64
        
        # Criar registro de pagamento pendente
        pagamento = Pagamento.objects.create(
            tipo_modelo='inscricao',
            content_type_id=8,
            object_id=inscricao.id,
            valor_pago=inscricao.valor_total,
            numero_parcela=1,
            status_pagamento='pendente',
            gateway_pagamento='pix_bancario',
            transaction_id=f"PIX_{inscricao.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}",
            payment_method='pix',
        )
        
        # Gerar dados do PIX
        pix_chave = settings.PIX_CHAVE
        pix_razao = settings.PIX_RAZAO_SOCIAL
        pix_cidade = settings.PIX_CIDADE
        
        # Se não tem chave PIX configurada, usar chave aleatória
        if not pix_chave:
            pix_chave = "pix@danca.com"
        
        if not pix_razao:
            pix_razao = "Sistema de Dança"
            
        if not pix_cidade:
            pix_cidade = "João Pessoa"
        
        # Montar EMV PIX (formato copia-e-cola)
        emv_pix = f"00020126580014BR.GOV.BCB.PIX0136{pix_chave}520400005303986540{float(inscricao.valor_total):.2f}5802BR5909{pix_razao[:25]:<25}6009{pix_cidade[:25]:<25}62170510***{pagamento.transaction_id}6304"
        
        # Calcular CRC16
        from zlib import crc32
        crc = format(crc32(emv_pix.encode()), 'x').upper()
        emv_pix += crc
        
        # Gerar QR Code
        qr = qrcode.QRCode(version=1, box_size=10, border=5)
        qr.add_data(emv_pix)
        qr.make(fit=True)
        
        img = qr.make_image(fill_color="black", back_color="white")
        
        # Converter para base64
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        qr_code_base64 = base64.b64encode(buffer.getvalue()).decode()
        
        return {
            'success': True,
            'pagamento_id': pagamento.id,
            'transaction_id': pagamento.transaction_id,
            'status': 'pendente',
            'payment_type': 'pix',
            'qr_code': emv_pix,
            'qr_code_base64': qr_code_base64,
            'chave_pix': pix_chave,
            'valor': str(inscricao.valor_total),
            'message': 'QR Code PIX gerado com sucesso. Escaneie o código para pagar.'
        }
    
    @staticmethod
    def _infinitepay_pix_payment(inscricao, dados_pix=None):
        """PIX direto de conta bancária - mesma lógica do PagSeguro"""
        return PaymentService._pagseguro_pix_payment(inscricao, dados_pix)

