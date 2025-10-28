from django.http import JsonResponse
from django.shortcuts import render, get_object_or_404
from django.urls import reverse, reverse_lazy
from django.views import View
from rest_framework import viewsets, permissions, filters
from .serializers import (
    CategoriaSerializer, EventoSerializer, CamisaSerializer, PlanejamentoSerializer,
    InscricaoSerializer, ProfissionalSerializer, EntradaSerializer, SaidaSerializer, 
    PagamentoSerializer, TipoEventoSerializer, LoteSerializer
)
from .models import Camisa, Planejamento, Inscricao, Profissional, Entrada, Saida, Pagamento, TipoEvento, Lote
from django.views.generic import ListView, TemplateView, DeleteView, DetailView
from django.contrib.contenttypes.models import ContentType
from django.contrib import messages
from django.utils.decorators import method_decorator
from django.utils import timezone
from django.views.decorators.cache import never_cache
from django.db.models.functions import Lower
from .models import Lote, Categoria, TipoEvento, Evento, Camisa,Planejamento,Inscricao, InscricaoEvento, Profissional, ProfissionalEvento, Entrada, Saida,Pagamento,PedidoCamisa
from .form import LoteForm,CategoriaForm,TipoEventoForm,EventoForm,CamisaForm,PlanejamentoForm,InscricaoForm, InscricaoEventoForm, ProfissionalForm, ProfissionalEventoForm, EntradaForm, SaidaForm, PagamentoForm,PedidoCamisaForm
from django.shortcuts import redirect,render
from django.db.models import Q
from django.core.exceptions import ValidationError
from django.db import models
from django.db.models import Sum, F, ExpressionWrapper, DecimalField, DateField,Value
from django.views.decorators.http import require_GET
from datetime import date
from django.db.models import OuterRef, Subquery
from django.contrib.contenttypes.models import ContentType
from django.db.models.functions import Coalesce
from django.views import View
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from .models import Inscricao
from django.utils.timezone import now
from django.db.models import Sum, F, Value, DecimalField, Q, Case, When
from django.db.models.functions import Coalesce



def index(request):
    return render(request, 'index.html')


#@method_decorator(login_required, name='dispatch')
@method_decorator(never_cache, name="dispatch")
class LoteListView(ListView):
    model = Lote


class CategoriaViewSet(viewsets.ModelViewSet):
    queryset = Categoria.objects.all().order_by("-id")
    serializer_class = CategoriaSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "descricao"]


class EventoViewSet(viewsets.ModelViewSet):
    queryset = Evento.objects.all().order_by("-id")
    serializer_class = EventoSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao", "tipo__descricao"]
    ordering_fields = ["id", "data", "descricao"]


class CamisaViewSet(viewsets.ModelViewSet):
    queryset = Camisa.objects.all().order_by("-id")
    serializer_class = CamisaSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "descricao"]


class PlanejamentoViewSet(viewsets.ModelViewSet):
    queryset = Planejamento.objects.all().order_by("-id")
    serializer_class = PlanejamentoSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "descricao"]


class InscricaoViewSet(viewsets.ModelViewSet):
    queryset = Inscricao.objects.all().order_by("-id")
    serializer_class = InscricaoSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["nome", "cpf"]
    ordering_fields = ["id", "nome"]
    
    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['request'] = self.request
        return context


class ProfissionalViewSet(viewsets.ModelViewSet):
    queryset = Profissional.objects.all().order_by("-id")
    serializer_class = ProfissionalSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["nome", "cpf"]
    ordering_fields = ["id", "nome"]
    
    def get_serializer_context(self):
        context = super().get_serializer_context()
        context['request'] = self.request
        return context


class EntradaViewSet(viewsets.ModelViewSet):
    queryset = Entrada.objects.all().order_by("-id")
    serializer_class = EntradaSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "data", "valor"]


class SaidaViewSet(viewsets.ModelViewSet):
    queryset = Saida.objects.all().order_by("-id")
    serializer_class = SaidaSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "data", "valor"]


class PagamentoViewSet(viewsets.ModelViewSet):
    queryset = Pagamento.objects.all().order_by("-id")
    serializer_class = PagamentoSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["observacoes"]
    ordering_fields = ["id", "data_pagamento", "valor_pago"]


class TipoEventoViewSet(viewsets.ModelViewSet):
    queryset = TipoEvento.objects.all().order_by("-id")
    serializer_class = TipoEventoSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "descricao"]


class LoteViewSet(viewsets.ModelViewSet):
    queryset = Lote.objects.all().order_by("-id")
    serializer_class = LoteSerializer
    permission_classes = [permissions.IsAuthenticatedOrReadOnly]
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ["descricao"]
    ordering_fields = ["id", "descricao"]


# Views antigas do Django (mantidas para compatibilidade)

#@method_decorator(login_required, name='dispatch')
@method_decorator(never_cache, name="dispatch")
class LoteDetailView(TemplateView):
    template_name = "lote/lote.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        lote = get_object_or_404(Lote, id=self.kwargs['lote_id'])
        context['lote'] = lote
        return context

#@method_decorator(login_required, name='dispatch')
@method_decorator(never_cache, name="dispatch")
class LoteFormView(View):
    form_class = LoteForm
    template_name = "lote/form.html"

    def get(self, request, lote_id=None):
        form = self.form_class()
        titulo = "Novo Lote" if not lote_id else "Editar Lote"
        if lote_id:
            lote = get_object_or_404(Lote, id=lote_id)
            form = self.form_class(instance=lote)
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
        })

    def post(self, request, lote_id=None):
        form = self.form_class(request.POST)
        msg = 'Lote criado com sucesso'

        if lote_id:
            lote = get_object_or_404(Lote, id=lote_id)
            form = self.form_class(request.POST, instance=lote)
            msg = 'Lote modificado com sucesso'
    
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_lotes')

#@method_decorator(login_required, name='dispatch')
@method_decorator(never_cache, name="dispatch")
class LoteDeleteView(DeleteView):  
    model = Lote
    pk_url_kwarg = "lote_id"    

    def get_success_url(self):
        messages.success(self.request, "Lote removido com sucesso")
        return reverse_lazy('list_lotes')

# CRUD de Categoria

@method_decorator(never_cache, name="dispatch")
class CategoriaListView(ListView):
    model = Categoria
    paginate_by = 10
    template_name = "categoria/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Categoria.objects.filter(descricao__icontains=query).order_by(Lower('descricao'))
        return Categoria.objects.all().order_by(Lower('descricao'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_categoria')  
        return context


@method_decorator(never_cache, name="dispatch")
class CategoriaDetailView(TemplateView):
    template_name = "categoria/categoria.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        categoria = get_object_or_404(Categoria, id=self.kwargs['categoria_id'])
        context['categoria'] = categoria
        return context

@method_decorator(never_cache, name="dispatch")
class CategoriaFormView(View):
    form_class = CategoriaForm
    template_name = "categoria/form.html"

    def get(self, request, categoria_id=None):
        form = self.form_class()
        titulo = "Nova Categoria" if not categoria_id else "Editar Categoria"        
        if categoria_id:
            categoria = get_object_or_404(Categoria, id=categoria_id)
            form = self.form_class(instance=categoria)
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,           
        })

    def post(self, request, categoria_id=None):
        form = self.form_class(request.POST)
        msg = 'Categoria criada com sucesso'

        if categoria_id:
            categoria = get_object_or_404(Categoria, id=categoria_id)
            form = self.form_class(request.POST, instance=categoria)
            msg = 'Categoria modificado com sucesso'
    
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_categorias')

#@method_decorator(login_required, name='dispatch')
@method_decorator(never_cache, name="dispatch")
class CategoriaDeleteView(DeleteView):  
    model = Categoria
    pk_url_kwarg = "categoria_id"    

    def get_success_url(self):
        messages.success(self.request, "Categoria removida com sucesso")
        return reverse_lazy('list_categorias')


@method_decorator(never_cache, name="dispatch")
class TipoEventoListView(ListView):
    model = TipoEvento
    paginate_by = 10
    template_name = "tipo_evento/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return TipoEvento.objects.filter(descricao__icontains=query).order_by(Lower('descricao'))
        return TipoEvento.objects.all().order_by(Lower('descricao'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_tipo_evento')  
        return context


@method_decorator(never_cache, name="dispatch")
class TipoEventoDetailView(TemplateView):
    template_name = "tipo_evento/tipo_evento.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        tipo_evento = get_object_or_404(TipoEvento, id=self.kwargs['tipo_evento_id'])
        context['tipo_evento'] = tipo_evento
        return context


@method_decorator(never_cache, name="dispatch")
class TipoEventoFormView(View):
    form_class = TipoEventoForm
    template_name = "tipo_evento/form.html"

    def get(self, request, tipo_evento_id=None):
        form = self.form_class()
        if tipo_evento_id:
            tipo_evento = get_object_or_404(TipoEvento, id=tipo_evento_id)
            form = self.form_class(instance=tipo_evento)
        return render(request, self.template_name, {"form": form})

    def post(self, request, tipo_evento_id=None):
        form = self.form_class(request.POST)
        msg = 'Tipo de Evento criado com sucesso'

        if tipo_evento_id:
            tipo_evento = get_object_or_404(TipoEvento, id=tipo_evento_id)
            form = self.form_class(request.POST, instance=tipo_evento)
            msg = 'Tipo de Evento modificado com sucesso'
    
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_tipo_eventos')


@method_decorator(never_cache, name="dispatch")
class TipoEventoDeleteView(DeleteView):  
    model = TipoEvento
    pk_url_kwarg = "tipo_evento_id"    

    def get_success_url(self):
        messages.success(self.request, "Tipo de Evento removido com sucesso")
        return reverse_lazy('list_tipo_eventos')


# CRUD de Evento

@method_decorator(never_cache, name="dispatch")
class EventoListView(ListView):
    model = Evento
    paginate_by = 10
    template_name = "evento/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Evento.objects.filter(descricao__icontains=query).order_by(Lower('descricao'))
        return Evento.objects.all().order_by(Lower('descricao'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_evento')  
        return context


@method_decorator(never_cache, name="dispatch")
class EventoDetailView(TemplateView):
    template_name = "evento/evento.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        evento = get_object_or_404(Evento, id=self.kwargs['evento_id'])
        context['evento'] = evento
        return context


@method_decorator(never_cache, name="dispatch")
class EventoFormView(View):
    form_class = EventoForm
    template_name = "evento/form.html"

    def get(self, request, evento_id=None):
        form = self.form_class()
        titulo = "Novo Evento" if not evento_id else "Editar Evento"
        if evento_id:
            evento = get_object_or_404(Evento, id=evento_id)
            form = self.form_class(instance=evento)
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
        })

    def post(self, request, evento_id=None):
        form = self.form_class(request.POST)
        msg = 'Evento criado com sucesso'

        if evento_id:
            evento = get_object_or_404(Evento, id=evento_id)
            form = self.form_class(request.POST, instance=evento)
            msg = 'Evento modificado com sucesso'
        
        if form.is_valid():
            evento = form.save()
            evento.atualizar_contador_inscricoes()
            messages.success(request, msg)
            return redirect('list_eventos')


@method_decorator(never_cache, name="dispatch")
class EventoDeleteView(DeleteView):  
    model = Evento
    pk_url_kwarg = "evento_id"    

    def get_success_url(self):
        messages.success(self.request, "Evento removido com sucesso")
        return reverse_lazy('list_eventos')

# CRUD Camisas
@method_decorator(never_cache, name="dispatch")
class CamisaListView(ListView):
    model = Camisa
    paginate_by = 10
    template_name = "camisas/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Camisa.objects.filter(descricao__icontains=query).order_by(Lower('descricao'))
        return Camisa.objects.all().order_by(Lower('descricao'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_camisa')
        return context

@method_decorator(never_cache, name="dispatch")
class CamisaDetailView(TemplateView):
    template_name = "camisas/camisa.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        camisa = get_object_or_404(Camisa, id=self.kwargs['camisa_id'])
        context['camisa'] = camisa
        context['pedidos'] = PedidoCamisa.objects.filter(camisa=camisa)
        return context

@method_decorator(never_cache, name="dispatch")
class CamisaFormView(View):
    form_class = CamisaForm
    template_name = "camisas/form.html"

    def get(self, request, camisa_id=None):
        form = self.form_class()
        titulo = "Nova Camisa"
        
        if camisa_id:
            camisa = get_object_or_404(Camisa, id=camisa_id)
            form = self.form_class(instance=camisa)
            titulo = "Editar Camisa"
        
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
        })

    def post(self, request, camisa_id=None):
        form = self.form_class(request.POST)
        msg = 'Camisa criada com sucesso'

        if camisa_id:
            camisa = get_object_or_404(Camisa, id=camisa_id)
            form = self.form_class(request.POST, instance=camisa)
            msg = 'Camisa atualizada com sucesso'
        
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_camisas')
        
        return render(request, self.template_name, {
            "form": form,
            "titulo": "Editar Camisa" if camisa_id else "Nova Camisa"
        })

@method_decorator(never_cache, name="dispatch")
class CamisaDeleteView(DeleteView):
    model = Camisa
    pk_url_kwarg = "camisa_id"

    def get_success_url(self):
        messages.success(self.request, "Camisa removida com sucesso")
        return reverse_lazy('list_camisas')

# CRUD Pedidos de Camisa
@method_decorator(never_cache, name="dispatch")
class PedidoCamisaListView(ListView):
    model = PedidoCamisa
    paginate_by = 10
    template_name = "pedidos/list.html"
    ordering = ['-data_pedido']

    def get_queryset(self):
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(nome_completo__icontains=query)
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_pedido')
        return context

@method_decorator(never_cache, name="dispatch")
class PedidoCamisaDetailView(TemplateView):
    template_name = "pedidos/detail.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        pedido = get_object_or_404(PedidoCamisa, id=self.kwargs['pedido_id'])
        context['pedido'] = pedido
        return context

@method_decorator(never_cache, name="dispatch")
class PedidoCamisaFormView(View):
    form_class = PedidoCamisaForm
    template_name = "pedidos/form.html"

    def get(self, request, pedido_id=None):
        form = self.form_class()
        titulo = "Novo Pedido"
        
        if pedido_id:
            pedido = get_object_or_404(PedidoCamisa, id=pedido_id)
            form = self.form_class(instance=pedido)
            titulo = "Editar Pedido"
        
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
        })

    def post(self, request, pedido_id=None):
        form = self.form_class(request.POST)
        msg = 'Pedido criado com sucesso'

        if pedido_id:
            pedido = get_object_or_404(PedidoCamisa, id=pedido_id)
            form = self.form_class(request.POST, instance=pedido)
            msg = 'Pedido atualizado com sucesso'
        
        if form.is_valid():
            pedido = form.save(commit=False)
            # Calcula o valor de venda automaticamente antes de salvar
            if pedido.camisa:
                # Se você quiser lógica diferente para sócios e não-sócios:
                if pedido.tipo_cliente == 'socio':
                    # Exemplo: desconto de 10% para sócios
                    pedido.valor_venda = pedido.camisa.valor_venda * Decimal('0.9')
                else:
                    # Preço normal para não-sócios
                    pedido.valor_venda = pedido.camisa.valor_venda
            pedido.save()
            messages.success(request, msg)
            return redirect('list_pedidos')
        
        return render(request, self.template_name, {
            "form": form,
            "titulo": "Editar Pedido" if pedido_id else "Novo Pedido"
        })

@method_decorator(never_cache, name="dispatch")
class PedidoCamisaDeleteView(DeleteView):
    model = PedidoCamisa
    pk_url_kwarg = "pedido_id"

    def get_success_url(self):
        messages.success(self.request, "Pedido removido com sucesso")
        return reverse_lazy('list_pedidos')
    
    
@method_decorator(never_cache, name="dispatch")   
class PlanejamentoListView(ListView):
    model = Planejamento
    paginate_by = 20
    template_name = "planejamento/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Planejamento.objects.filter(descricao__icontains=query).order_by(Lower('descricao'))
        return Planejamento.objects.all().order_by(Lower('descricao'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['total_planejado'] = Planejamento.objects.aggregate(total=Sum('valor_planejado'))['total'] or 0
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_planejamento')
        return context


@method_decorator(never_cache, name="dispatch")
class PlanejamentoDetailView(TemplateView):
    template_name = "planejamento/planejamento.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        planejamento = get_object_or_404(Planejamento, id=self.kwargs['planejamento_id'])
        context['planejamento'] = planejamento
        return context


@method_decorator(never_cache, name="dispatch")
class PlanejamentoFormView(View):
    form_class = PlanejamentoForm
    template_name = "planejamento/form.html"

    def get(self, request, planejamento_id=None):
        form = self.form_class()
        titulo = "Novo Planejamento" if not planejamento_id else "Editar Planejamento"
        if planejamento_id:
            planejamento = get_object_or_404(Planejamento, id=planejamento_id)
            form = self.form_class(instance=planejamento)
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
        })

    def post(self, request, planejamento_id=None):
        form = self.form_class(request.POST)
        msg = 'Planejamento criado com sucesso'

        if planejamento_id:
            planejamento = get_object_or_404(Planejamento, id=planejamento_id)
            form = self.form_class(request.POST, instance=planejamento)
            msg = 'Planejamento modificado com sucesso'
    
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_planejamentos')


@method_decorator(never_cache, name="dispatch")
class PlanejamentoDeleteView(DeleteView):
    model = Planejamento
    pk_url_kwarg = "planejamento_id"

    def get_success_url(self):
        messages.success(self.request, "Planejamento removido com sucesso")
        return reverse_lazy('list_planejamentos')


@method_decorator(never_cache, name="dispatch")
class InscricaoListView(ListView):
    model = Inscricao
    paginate_by = 50
    template_name = "inscricao/list.html"

    def get_queryset(self):
        ordering = self.request.GET.get('ordering', 'nome')  # Ordena por 'nome' por padrão
        filtro = self.request.GET.get('q', '').strip().lower()  # Obtém o parâmetro de busca 'q'
        proximo_pagamento_filter = self.request.GET.get('proximo_pagamento_filter')
        status_filter = self.request.GET.get('status_filter')  
        uf_filter = self.request.GET.get('uf_filter')  
        categoria_filter = self.request.GET.get('categoria_filter')   

        content_type = ContentType.objects.get_for_model(Inscricao)

        inscricoes = Inscricao.objects.annotate(
            valor_pago_db=Coalesce(
                Subquery(
                    Pagamento.objects.filter(
                        content_type=content_type,
                        object_id=OuterRef('id')
                    ).values('object_id')
                    .annotate(total=Sum('valor_pago'))
                    .values('total')[:1]  # Adicionado [:1] para garantir um único resultado
                ),
                Value(0),
                output_field=models.DecimalField()
            ),
            valor_restante_db=ExpressionWrapper(
                F('valor_total') - F('valor_pago_db'),
                output_field=models.DecimalField()
            )
        )

        # Filtro de busca por nome (como no PagamentoListView)
        if filtro:
            inscricoes = inscricoes.filter(nome__icontains=filtro)
            
        # filtro por UF    
        if uf_filter:
            inscricoes = inscricoes.filter(uf=uf_filter)
        
        #filtro por categoria
        if categoria_filter:
            inscricoes = inscricoes.filter(categoria_id=categoria_filter)

         # Subquery para próximo pagamento (atualizada)
        hoje = timezone.now().date()

        proximo_pagamento_subquery = Pagamento.objects.filter(
            content_type=ContentType.objects.get_for_model(Inscricao),
            object_id=OuterRef('id'),
            data_proximo_pagamento__gte=hoje
        ).order_by('data_proximo_pagamento').values('data_proximo_pagamento')[:1]

        ultimo_pagamento_subquery = Pagamento.objects.filter(
            content_type=ContentType.objects.get_for_model(Inscricao),
            object_id=OuterRef('id')
        ).order_by('-data_proximo_pagamento').values('data_proximo_pagamento')[:1]

        inscricoes = inscricoes.annotate(
            proximo_pagamento_futuro=Subquery(proximo_pagamento_subquery),
            ultimo_pagamento=Subquery(ultimo_pagamento_subquery)
        ).annotate(
            data_proximo_pagamento=Coalesce(
                'proximo_pagamento_futuro',
                'ultimo_pagamento',
                output_field=DateField()
            )
        )

        # Filtro por status do próximo pagamento
        if proximo_pagamento_filter:
            if proximo_pagamento_filter == 'atrasado':
                inscricoes = inscricoes.filter(data_proximo_pagamento__lt=hoje)
            elif proximo_pagamento_filter == 'hoje':
                inscricoes = inscricoes.filter(data_proximo_pagamento=hoje)
            elif proximo_pagamento_filter == 'futuro':
                inscricoes = inscricoes.filter(data_proximo_pagamento__gt=hoje)

         # Filtro por Status (usando annotation)
        if status_filter:
            if status_filter == 'pago':
                inscricoes = inscricoes.filter(valor_restante_db__lte=0)  # Alterado
            elif status_filter == 'parcial':
                inscricoes = inscricoes.filter(valor_pago_db__gt=0, valor_restante_db__gt=0)  # Alterado
            elif status_filter == 'pendente':
                inscricoes = inscricoes.filter(valor_pago_db=0)
                    
        # Mapa de ordenações permitidas
        ordering_map = {
            'nome': 'nome',
            '-nome': '-nome',
            'categoria': 'categoria__descricao',
            '-categoria': '-categoria__descricao',
            'uf': 'uf',  
            '-uf': '-uf',
            'lote': 'lote__descricao',
            '-lote': '-lote__descricao',
            'valor_total': 'valor_total',
            '-valor_total': '-valor_total',
            'numero_parcelas': 'numero_parcelas',
            '-numero_parcelas': '-numero_parcelas',
            'valor_parcela': 'valor_parcela',
            '-valor_parcela': '-valor_parcela',
            'valor_restante': 'valor_restante_db',
            '-valor_restante': '-valor_restante_db',
            'proximo_pagamento': 'data_proximo_pagamento',  
            '-proximo_pagamento': '-data_proximo_pagamento',
        }

        if ordering in ordering_map:
            return inscricoes.order_by(ordering_map[ordering])
        return inscricoes.order_by('nome')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        queryset = self.get_queryset()
        
        # Cálculo dos totais (mantido do seu código original)
        total_pago = sum(inscricao.valor_pago for inscricao in queryset)
        total_total = queryset.aggregate(total=Sum('valor_total'))['total'] or 0
        total_a_receber = total_total - total_pago

            # Busca as UFs existentes
        ufs_existentes = Inscricao.objects.order_by('uf').values_list('uf', flat=True).distinct()
        UFS_BRASIL = [...]  # Sua lista de UFs
        
        # Busca apenas as categorias que estão sendo usadas em inscrições
        categorias_em_uso = Categoria.objects.filter(
            inscricao__isnull=False
        ).distinct().order_by('descricao')

         # Lista de UFs do Brasil para o filtro
        UFS_BRASIL = [
            ('AC', 'Acre'), ('AL', 'Alagoas'), ('AP', 'Amapá'),
            ('AM', 'Amazonas'), ('BA', 'Bahia'), ('CE', 'Ceará'),
            ('DF', 'Distrito Federal'), ('ES', 'Espírito Santo'),
            ('GO', 'Goiás'), ('MA', 'Maranhão'), ('MT', 'Mato Grosso'),
            ('MS', 'Mato Grosso do Sul'), ('MG', 'Minas Gerais'),
            ('PA', 'Pará'), ('PB', 'Paraíba'), ('PR', 'Paraná'),
            ('PE', 'Pernambuco'), ('PI', 'Piauí'), ('RJ', 'Rio de Janeiro'),
            ('RN', 'Rio Grande do Norte'), ('RS', 'Rio Grande do Sul'),
            ('RO', 'Rondônia'), ('RR', 'Roraima'), ('SC', 'Santa Catarina'),
            ('SP', 'São Paulo'), ('SE', 'Sergipe'), ('TO', 'Tocantins')
        ]
        # Contexto (simplificado, apenas com o necessário)
        context.update({
            'q': self.request.GET.get('q', ''),  # Adiciona o valor do filtro ao contexto
            'ordering': self.request.GET.get('ordering', 'nome'),
            'proximo_pagamento_filter': self.request.GET.get('proximo_pagamento_filter', ''),
            'status_filter': self.request.GET.get('status_filter', ''),
            'uf_filter': self.request.GET.get('uf_filter', ''),  # Novo
            'categoria_filter': self.request.GET.get('categoria_filter', ''),
            'categorias_disponiveis': categorias_em_uso,
            'ufs_brasil': UFS_BRASIL,  # Lista de UFs para o template
            'create_url': reverse('create_inscricao'),
            'total_pago': total_pago,
            'total_a_receber': total_a_receber,
            'total_inscricoes': Inscricao.objects.count(),
            'search_url': self.request.path,  # Importante para manter a paginação com filtros
        })
        return context

@method_decorator(never_cache, name="dispatch")
class InscricaoDetailView(TemplateView):
    template_name = "inscricao/incricao.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        inscricao = get_object_or_404(Inscricao, id=self.kwargs['inscricao_id'])
        context['inscricao'] = inscricao
        return context


@method_decorator(never_cache, name="dispatch")
class InscricaoFormView(View):
    form_class = InscricaoForm
    template_name = "inscricao/form.html"

    def get(self, request, inscricao_id=None):
        form = self.form_class()
        eventos = Evento.objects.all()
        inscricao = None

        if inscricao_id:
            inscricao = get_object_or_404(Inscricao, id=inscricao_id)
            form = self.form_class(instance=inscricao)

        # Carrega a data_proximo_pagamento do último pagamento se existir
            ultimo_pagamento = Pagamento.objects.filter(
                content_type=ContentType.objects.get_for_model(inscricao),
                object_id=inscricao.id,
                tipo_modelo='inscricao'
            ).order_by('-data_pagamento').first()
            
            if ultimo_pagamento:
                form.fields['data_proximo_pagamento'].initial = ultimo_pagamento.data_proximo_pagamento

        # Certifique-se de que o queryset do lote está carregado
        form.fields['lote'].queryset = Lote.objects.all()

        return render(request, self.template_name, {
            "form": form,
            "eventos": eventos,
            "inscricao": inscricao,
        })

    def post(self, request, inscricao_id=None):
        inscricao = None
        if inscricao_id:
            inscricao = get_object_or_404(Inscricao, id=inscricao_id)
            form = self.form_class(request.POST, instance=inscricao)
        else:
            form = self.form_class(request.POST)

        if form.is_valid():
            inscricao = form.save(commit=False)
            inscricao.save()

            # Associa os eventos selecionados
            eventos_ids = request.POST.getlist('eventos')
            eventos = Evento.objects.filter(id__in=eventos_ids)
            
            # Verifica se há vagas suficientes para cada evento
            for evento in eventos:
                if evento.quantidade_pessoas is not None:
                    if evento.quantidade_pessoas <= 0:
                        raise ValidationError(f"O evento '{evento.descricao}' não tem vagas disponíveis.")
                    
                    

            inscricao.eventos.set(eventos)

            # Atualiza o valor total e o valor da parcela
            inscricao.valor_total = inscricao.calcular_valor_total()
            inscricao.valor_parcela = inscricao.calcular_valor_parcela()
            inscricao.save(update_fields=['valor_total', 'valor_parcela'])

            # Atualização MANUAL da data de próximo pagamento (apenas para edição)
            if inscricao_id:
                data_proximo = form.cleaned_data.get('data_proximo_pagamento')
                if data_proximo:
                    # Busca o próximo pagamento não pago ou cria um novo se necessário
                    pagamento = Pagamento.objects.filter(
                        content_type=ContentType.objects.get_for_model(inscricao),
                        object_id=inscricao.id,
                        tipo_modelo='inscricao',
                        valor_pago=0  # Ainda não foi pago
                    ).order_by('numero_parcela').first()
                    
                    if not pagamento:
                        # Se não encontrou, pega o último pagamento
                        pagamento = Pagamento.objects.filter(
                            content_type=ContentType.objects.get_for_model(inscricao),
                            object_id=inscricao.id,
                            tipo_modelo='inscricao'
                        ).order_by('-numero_parcela').first()
                    
                    if pagamento:
                        pagamento.data_proximo_pagamento = data_proximo
                        pagamento.save(update_fields=['data_proximo_pagamento'])

            messages.success(request, "Inscrição salva com sucesso!")
            return redirect('list_inscricoes')
            
        # Garante que o lote esteja carregado mesmo em caso de erro
        form.fields['lote'].queryset = Lote.objects.all()    

        return render(request, self.template_name, {
            "form": form,
            "eventos": Evento.objects.all(),
            "inscricao": inscricao,
        })

@method_decorator(never_cache, name="dispatch")
class InscricaoDeleteView(DeleteView):
    model = Inscricao
    pk_url_kwarg = "inscricao_id"

    def get_success_url(self):
        messages.success(self.request, "Inscrição removida com sucesso")
        return reverse_lazy('list_inscricoes')

@method_decorator(never_cache, name="dispatch")
class InscricaoEventoFormView(View):
    form_class = InscricaoEventoForm
    template_name = "inscricao/form_evento.html"

    def get(self, request, inscricao_id):
        form = self.form_class()
        eventos = Evento.objects.all()
        inscricao = get_object_or_404(Inscricao, id=inscricao_id)
        eventos = inscricao.eventos.all()  # Recupera os eventos associados      
            
        return render(request, self.template_name, {
            "form": form,
            "inscricao": inscricao,
            'eventos': eventos,
        })

    def post(self, request, inscricao_id):
        form = self.form_class(request.POST)
        inscricao = get_object_or_404(Inscricao, id=inscricao_id)
        msg = 'Evento adicionado à inscrição com sucesso'

        if form.is_valid():
            inscricao = form.save(commit=False)
            
            # Ainda não salva no banco
            eventos_ids = request.POST.getlist('eventos')
            eventos = Evento.objects.filter(id__in=eventos_ids)
            
            inscricao.valor_total = sum(evento.valor_unitario for evento in eventos) - inscricao.desconto + inscricao.lote.valor_unitario
            inscricao.valor_parcela = inscricao.calcular_valor_parcela()

            inscricao.save()  # Agora sim, já com valores certos
            inscricao.eventos.set(eventos)

            # Atualiza o contador de inscrições
            for evento in eventos:
                evento.atualizar_contador_inscricoes()

            return redirect('list_inscricoes')

@method_decorator(never_cache, name="dispatch")
class InscricaoEventoDeleteView(DeleteView):
    model = InscricaoEvento
    pk_url_kwarg = "inscricao_evento_id"

    def get_success_url(self):
        inscricao_id = self.kwargs['inscricao_id']
        inscricao = get_object_or_404(Inscricao, id=inscricao_id)
        
        # Recalcular valor total da inscrição após remover o evento
        inscricao.valor_total = inscricao.calcular_valor_total()
        inscricao.save()
        
        # Atualiza o contador de inscrições para o evento associado
        inscricao_evento = self.object
        inscricao_evento.evento.atualizar_contador_inscricoes()
        
        messages.success(self.request, "Evento removido da inscrição com sucesso")
        return reverse('detail_inscricao', kwargs={'inscricao_id': inscricao_id})

@method_decorator(never_cache, name="dispatch")
class InscricaoCreateView(View):
    def get(self, request):
        form = InscricaoForm()
        return render(request, 'inscricao_form.html', {'form': form})

    def post(self, request):
        form = InscricaoForm(request.POST)
        if form.is_valid():
            # Salva a inscrição sem os eventos
            inscricao = form.save(commit=False)
            inscricao.save()
            

            # Associa os eventos selecionados
            eventos = form.cleaned_data['eventos']
            inscricao.eventos.set(eventos)

            # Atualiza o valor total e o valor da parcela
            inscricao.valor_total = inscricao.calcular_valor_total()
            inscricao.valor_parcela = inscricao.calcular_valor_parcela()
            inscricao.save(update_fields=['valor_total', 'valor_parcela'])

            return redirect('list_inscricoes')  # Redireciona para a lista de inscrições
        return render(request, 'inscricao_form.html', {'form': form})

@method_decorator(never_cache, name="dispatch")
class EventoInscritosView(DetailView):
    model = Evento
    template_name = "evento/inscritos.html"
    context_object_name = "evento"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Congressistas (antigos "inscritos")
        inscritos = (
            InscricaoEvento.objects
            .filter(evento=self.object)
            .select_related('inscricao')
            .order_by('inscricao__nome')
        )

        lista_congressistas = [
            {'nome': i.inscricao.nome, 'cpf': i.inscricao.cpf}
            for i in inscritos
        ]

        # Profissionais
        profissionais = (
            ProfissionalEvento.objects
            .filter(evento=self.object)
            .select_related('profissional')
            .order_by('profissional__nome')
        )

        lista_profissionais = [
            {'nome': p.profissional.nome, 'cpf': p.profissional.cpf}
            for p in profissionais
        ]

        total_congressistas = len(lista_congressistas)
        total_profissionais = len(lista_profissionais)
        total_geral = total_congressistas + total_profissionais

        context.update({
            'congressistas': lista_congressistas,
            'profissionais': lista_profissionais,
            'total_congressistas': total_congressistas,
            'total_profissionais': total_profissionais,
            'total_geral': total_geral,
            'capacidade': 130,  # opcional
        })
        return context


@method_decorator(never_cache, name="dispatch")
class ProfissionalListView(ListView):
    model = Profissional
    paginate_by = 30
    template_name = "profissional/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Profissional.objects.filter(nome__icontains=query).order_by(Lower('nome'))
        return Profissional.objects.all().order_by(Lower('nome'))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '')
        context['create_url'] = reverse('create_profissional')
        return context


@method_decorator(never_cache, name="dispatch")
class ProfissionalDetailView(TemplateView):
    template_name = "profissional/profissional.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        profissional = get_object_or_404(Profissional, id=self.kwargs['profissional_id'])
        context['profissional'] = profissional
        return context


@method_decorator(never_cache, name="dispatch")
class ProfissionalFormView(View):
    form_class = ProfissionalForm
    template_name = "profissional/form.html"

    def get(self, request, profissional_id=None):
        form = self.form_class()
        profissional = None
        if profissional_id:
            profissional = get_object_or_404(Profissional, id=profissional_id)
            form = self.form_class(instance=profissional)
        eventos = Evento.objects.all()
        return render(request, self.template_name, {
            "form": form,
            "profissional": profissional,
            "eventos": eventos,
            "titulo": "Editar Profissional" if profissional else "Novo Profissional"
        })

    def post(self, request, profissional_id=None):
        profissional = None
        if profissional_id:
            profissional = get_object_or_404(Profissional, id=profissional_id)
            form = self.form_class(request.POST, instance=profissional)
        else:
            form = self.form_class(request.POST)

        if form.is_valid():
            profissional = form.save(commit=False)
            profissional.save()

            # ⭐ EVENTOS SÃO OPCIONAIS - sempre limpa e adiciona os selecionados
            eventos_ids = request.POST.getlist('eventos')
            eventos_validos = Evento.objects.filter(id__in=eventos_ids)
            profissional.eventos.set(eventos_validos)  # Se vazio, limpa todos

            messages.success(request, "Profissional salvo com sucesso!")
            return redirect('list_profissionais')

        return render(request, self.template_name, {
            "form": form,
            "profissional": profissional,
            "eventos": Evento.objects.all(),
            "titulo": "Editar Profissional" if profissional else "Novo Profissional"
        })
@method_decorator(never_cache, name="dispatch")
class ProfissionalDeleteView(DeleteView):
    model = Profissional
    pk_url_kwarg = "profissional_id"

    def get_success_url(self):
        messages.success(self.request, "Profissional removido com sucesso")
        return reverse_lazy('list_profissionais')


@method_decorator(never_cache, name="dispatch")
class ProfissionalEventoFormView(View):
    form_class = ProfissionalEventoForm
    template_name = "profissional/form_evento.html"

    def get(self, request, profissional_id):
        form = self.form_class()
        eventos = Evento.objects.all()
        profissional = get_object_or_404(Profissional, id=profissional_id)
        eventos = profissional.eventos.all()
        
        return render(request, self.template_name, {
            "form": form,
            "profissional": profissional,
            'eventos': eventos,
        })

    def post(self, request, profissional_id):
        form = self.form_class(request.POST)
        profissional = get_object_or_404(Profissional, id=profissional_id)
        msg = 'Evento adicionado ao profissional com sucesso'

        if form.is_valid():
            profissional_evento = form.save(commit=False)
            profissional_evento.profissional = profissional
            profissional_evento.save()

            # Atualiza o contador de inscrições para o evento
            profissional_evento.evento.atualizar_contador_inscricoes()

            messages.success(request, msg)
            return redirect('list_profissionais')


@method_decorator(never_cache, name="dispatch")
class ProfissionalEventoDeleteView(DeleteView):
    model = ProfissionalEvento
    pk_url_kwarg = "profissional_evento_id"

    def get_success_url(self):
        profissional_id = self.kwargs['profissional_id']
        profissional = get_object_or_404(Profissional, id=profissional_id)
        
        # Atualiza o contador de inscrições para o evento associado
        profissional_evento = self.object
        profissional_evento.evento.atualizar_contador_inscricoes()
        
        messages.success(self.request, "Evento removido do profissional com sucesso")
        return reverse('detail_profissional', kwargs={'profissional_id': profissional_id})

@method_decorator(never_cache, name="dispatch")
class EntradaListView(ListView):
    model = Entrada
    paginate_by = 10
    template_name = "entrada/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Entrada.objects.filter(descricao__icontains=query).order_by('-data')
        return Entrada.objects.all().order_by('-data')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['total_entrada'] = Entrada.objects.aggregate(total=Sum('valor'))['total'] or 0
        context['q'] = self.request.GET.get('q', '')
        context['total_entradas'] = sum(entrada.valor for entrada in context['object_list'])
        return context

@method_decorator(never_cache, name="dispatch")
class EntradaDetailView(TemplateView):
    template_name = "entrada/entrada.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        entrada = get_object_or_404(Entrada, id=self.kwargs['entrada_id'])
        context['entrada'] = entrada
        return context

@method_decorator(never_cache, name="dispatch")
class EntradaFormView(View):
    form_class = EntradaForm
    template_name = "entrada/form.html"

    def get(self, request, entrada_id=None):
        form = self.form_class()
        titulo = "Nova Entrada" if not entrada_id else "Editar Entrada"
        if entrada_id:
            entrada = get_object_or_404(Entrada, id=entrada_id)
            form = self.form_class(instance=entrada)
        return render(request, self.template_name, {"form": form,"titulo": titulo,})

    def post(self, request, entrada_id=None):
        form = self.form_class(request.POST)
        msg = 'Entrada criada com sucesso'

        if entrada_id:
            entrada = get_object_or_404(Entrada, id=entrada_id)
            form = self.form_class(request.POST, instance=entrada)
            msg = 'Entrada modificada com sucesso'
    
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_entradas')

@method_decorator(never_cache, name="dispatch")
class EntradaDeleteView(DeleteView):  
    model = Entrada
    pk_url_kwarg = "entrada_id"    

    def get_success_url(self):
        messages.success(self.request, "Entrada removida com sucesso")
        return reverse_lazy('list_entradas')

@method_decorator(never_cache, name="dispatch")
class SaidaListView(ListView):
    model = Saida
    paginate_by = 10
    template_name = "saida/list.html"

    def get_queryset(self):
        query = self.request.GET.get('q')
        if query:
            return Saida.objects.filter(descricao__icontains=query).order_by('-data') # type: ignore
        return Saida.objects.all().order_by('-data')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['total_saida'] = Saida.objects.aggregate(total=Sum('valor'))['total'] or 0
        context['q'] = self.request.GET.get('q', '')
        context['total_saidas'] = sum(saida.valor for saida in context['object_list'])
        return context

@method_decorator(never_cache, name="dispatch")
class SaidaDetailView(TemplateView):
    template_name = "saida/saida.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        saida = get_object_or_404(Saida, id=self.kwargs['saida_id'])
        context['saida'] = saida
        return context

@method_decorator(never_cache, name="dispatch")
class SaidaFormView(View):
    form_class = SaidaForm
    template_name = "saida/form.html"

    def get(self, request, saida_id=None):
        form = self.form_class()
        titulo = "Nova Saída" if not saida_id else "Editar Saída"
        if saida_id:
            saida = get_object_or_404(Saida, id=saida_id)
            form = self.form_class(instance=saida)
        return render(request, self.template_name, {"form": form,"titulo": titulo,})

    def post(self, request, saida_id=None):
        form = self.form_class(request.POST)
        msg = 'Saída criada com sucesso'

        if saida_id:
            saida = get_object_or_404(Saida, id=saida_id)
            form = self.form_class(request.POST, instance=saida)
            msg = 'Saída modificada com sucesso'
    
        if form.is_valid():
            form.save()
            messages.success(request, msg)
            return redirect('list_saidas')

@method_decorator(never_cache, name="dispatch")
class SaidaDeleteView(DeleteView):  
    model = Saida
    pk_url_kwarg = "saida_id"    

    def get_success_url(self):
        messages.success(self.request, "Saída removida com sucesso")
        return reverse_lazy('list_saidas')

def resumo_caixa(request):
    # Entradas e saídas
    total_entradas = Entrada.objects.aggregate(total=Sum('valor'))['total'] or 0
    total_saidas = Saida.objects.aggregate(total=Sum('valor'))['total'] or 0

    # Total inscricoes pagas (somando os pagamentos com tipo_modelo='inscricao')
    ct_inscricao = ContentType.objects.get_for_model(Inscricao)
    total_pago_inscricoes = Pagamento.objects.filter(content_type=ct_inscricao).aggregate(total=Sum('valor_pago'))['total'] or 0

    # Total a receber de inscrições (valor_total - valor_pago)
    total_valor_inscricoes = Inscricao.objects.aggregate(total=Sum('valor_total'))['total'] or 0
    total_a_receber = total_valor_inscricoes - total_pago_inscricoes

    # Total camisas - SOMENTE pedidos com status 'pago' ou 'entregue'
    total_camisas_pagas = (
        PedidoCamisa.objects
        .filter(Q(status='pago') | Q(status='entregue'))
        .aggregate(
            total=Coalesce(
                Sum(
                    Case(
                        When(tipo_cliente='equipe', then=Value(0, output_field=DecimalField())),
                        When(tipo_cliente='colaborador', then=F('camisa__valor_compra')),
                        default=F('valor_venda'),
                        output_field=DecimalField()
                    )
                ),
                Value(0, output_field=DecimalField())
            )
        )['total']
    )

    # Total planejado e total pago em planejamentos
    total_planejamentos = Planejamento.objects.aggregate(total=Sum('valor_planejado'))['total'] or 0
    ct_planejamento = ContentType.objects.get_for_model(Planejamento)
    total_pago_planejamento = Pagamento.objects.filter(content_type=ct_planejamento).aggregate(total=Sum('valor_pago'))['total'] or 0

    # Valor a pagar = planejado - pago
    total_a_pagar = total_planejamentos - total_pago_planejamento

    # Saldo em caixa = entradas + inscrições pagas + camisas - saídas - pagamentos de planejamento
    saldo_caixa = (total_entradas + total_pago_inscricoes + total_camisas_pagas) - (total_saidas + total_pago_planejamento)

     # Cálculo da estimativa
    saldo_futuro_previsto = (saldo_caixa + total_a_receber) - (total_planejamentos - total_saidas)

    context = {
        'total_entradas': total_entradas,
        'total_saidas': total_saidas,
        'total_inscricoes': total_pago_inscricoes,  # Total recebido das inscrições
        'total_a_receber': total_a_receber,
        'total_planejamentos': total_planejamentos,
        'total_pago_planejamento': total_pago_planejamento,
        'total_a_pagar': total_a_pagar,
        'total_camisas_pagas': total_camisas_pagas,
        'saldo_caixa': saldo_caixa,
        'saldo_futuro_previsto': saldo_futuro_previsto,
    }

    return render(request, 'resumo/resumo_caixa.html', context)

# Lista de Pagamentos
@method_decorator(never_cache, name="dispatch")
class PagamentoListView(ListView):
    model = Pagamento
    paginate_by = 35
    template_name = "pagamento/list.html"

    def get_queryset(self):
        queryset = Pagamento.objects.all().order_by('-id')
        filtro = self.request.GET.get('q', '').strip().lower()

        if filtro:
            # Busca os objetos relacionados que contenham o texto
            planejamentos = Planejamento.objects.filter(descricao__icontains=filtro)
            inscricoes = Inscricao.objects.filter(nome__icontains=filtro)

            # Monta a lista de conteúdo genérico
            content_type_planejamento = ContentType.objects.get_for_model(Planejamento)
            content_type_inscricao = ContentType.objects.get_for_model(Inscricao)

            queryset = queryset.filter(
                models.Q(content_type=content_type_planejamento, object_id__in=planejamentos.values_list('id', flat=True)) |
                models.Q(content_type=content_type_inscricao, object_id__in=inscricoes.values_list('id', flat=True))
            )

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        filtro = self.request.GET.get('q', '')
        context['pagamento_relacionado'] = filtro
        context['search_url'] = self.request.path
        context['create_url'] = reverse('create_pagamento')
        return context
           
# Criar e Editar Pagamento

@method_decorator(never_cache, name="dispatch")
class PagamentoFormView(View):
    form_class = PagamentoForm
    template_name = "pagamento/form.html"

    def get(self, request, pagamento_id=None):
        if pagamento_id:
            pagamento = get_object_or_404(Pagamento, pk=pagamento_id)
            form = self.form_class(instance=pagamento)
            titulo = "Editar Pagamento"
        else:
            pagamento = None
            form = self.form_class()
            titulo = "Novo Pagamento"
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
            "pagamento": pagamento,
        })

    def post(self, request, pagamento_id=None):
        if pagamento_id:
            pagamento = get_object_or_404(Pagamento, pk=pagamento_id)
            form = self.form_class(request.POST, instance=pagamento)
        else:
            pagamento = None
            form = self.form_class(request.POST)

        if form.is_valid():
            form.save()
            messages.success(request, "Pagamento salvo com sucesso!")
            return redirect('list_pagamentos')
        else:
            print("Formulário inválido:", form.errors)

        titulo = "Editar Pagamento" if pagamento_id else "Novo Pagamento"
        return render(request, self.template_name, {
            "form": form,
            "titulo": titulo,
            "pagamento": pagamento,
        })


# Detalhamento do Pagamento
@method_decorator(never_cache, name="dispatch")
class PagamentoDetailView(TemplateView):
    template_name = "pagamento/pagamento.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        pagamento = get_object_or_404(Pagamento, id=self.kwargs['pagamento_id'])
        context['pagamento'] = pagamento
        return context


# Remover Pagamento
@method_decorator(never_cache, name="dispatch")
class PagamentoDeleteView(DeleteView):
    model = Pagamento
    pk_url_kwarg = "pagamento_id"

    def get_success_url(self):
        messages.success(self.request, "Pagamento removido com sucesso")
        return reverse_lazy('list_pagamentos')


# Carregar objetos relacionados com AJAX
@require_GET
def carregar_objetos_pagamento(request):
    tipo = request.GET.get('tipo_modelo')
    termo = request.GET.get('q', '')

    if tipo == 'planejamento':
        objetos = Planejamento.objects.filter(descricao__icontains=termo)
    elif tipo == 'inscricao':
        objetos = Inscricao.objects.filter(nome__icontains=termo)
    else:
        objetos = []

    data = {
        'objetos': [{'id': obj.id, 'descricao': str(obj)} for obj in objetos]
    }
    return JsonResponse(data)

def pagamentos_list(request):
    pagamentos = Pagamento.objects.all()
    tipo_modelo = request.GET.get('tipo_modelo')
    pagamento_relacionado = request.GET.get('pagamento_relacionado')

    if tipo_modelo:
        pagamentos = pagamentos.filter(tipo_modelo=tipo_modelo)

    if pagamento_relacionado:
        pagamentos = pagamentos.filter(pagamento_relacionado_id=pagamento_relacionado)

    # Popula opções para pagamento_relacionado conforme tipo_modelo
    if tipo_modelo == 'planejamento':
        pagamentos_relacionados = Planejamento.objects.all()
    elif tipo_modelo == 'inscricao':
        pagamentos_relacionados = Inscricao.objects.all()
    else:
        pagamentos_relacionados = []

    context = {
        'page_obj': pagamentos,
        'pagamentos_relacionados': pagamentos_relacionados,
        'create_url': reverse('create_pagamento'),
    }
    return render(request, 'pagamentos.html', context)


def listar_pagamentos(request):
    pagamentos = Pagamento.objects.all()

    tipo_modelo = request.GET.get('tipo_modelo')
    pagamento_relacionado = request.GET.get('pagamento_relacionado')

    if tipo_modelo:
        pagamentos = pagamentos.filter(tipo_modelo=tipo_modelo)

    if pagamento_relacionado:
        pagamentos = pagamentos.filter(pagamento_relacionado__id=pagamento_relacionado)

    paginator = Paginator(pagamentos, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    return render(request, 'pagamentos/lista.html', {
        'page_obj': page_obj,
    })

class InscricaoRelatorioDocxView(InscricaoListView):
    """ Gera relatório .docx como lista enumerada """
    def get(self, request, *args, **kwargs):
        queryset = self.get_queryset()

        document = Document()
        document.add_heading('Relatório de Inscrições', 0)
        document.add_paragraph(f'Gerado em: {now().strftime("%d/%m/%Y %H:%M")}')
        document.add_paragraph(f'Total de inscrições: {queryset.count()}')
        document.add_paragraph('')

        for inscricao in queryset:
            categoria = inscricao.categoria.descricao if inscricao.categoria else 'Sem categoria'
            if inscricao.valor_restante_db <= 0:
                status = 'Pago'
            elif inscricao.valor_pago_db > 0:
                status = 'Parcial'
            else:
                status = 'Pendente'
            document.add_paragraph(
                f'{inscricao.nome} — {categoria} — {status}',
                style='List Number'
            )


        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="relatorio_inscricoes.docx"'
        document.save(response)
        return response



def evento_inscritos_docx(request, pk):
    evento = Evento.objects.get(pk=pk)

    # Congressistas
    inscritos = (
        InscricaoEvento.objects
        .filter(evento=evento)
        .select_related('inscricao')
        .order_by('inscricao__nome')
    )
    lista_congressistas = [
        {
            'nome': i.inscricao.nome or 'Sem Cadastro',
            'cpf': i.inscricao.cpf or 'Sem Cadastro'
        }
        for i in inscritos
    ]

    # Profissionais
    profissionais = (
        ProfissionalEvento.objects
        .filter(evento=evento)
        .select_related('profissional')
        .order_by('profissional__nome')
    )
    lista_profissionais = [
        {
            'nome': p.profissional.nome or 'Sem Cadastro',
            'cpf': p.profissional.cpf or 'Sem Cadastro'
        }
        for p in profissionais
    ]

    total_congressistas = len(lista_congressistas)
    total_profissionais = len(lista_profissionais)
    total_geral = total_congressistas + total_profissionais

    document = Document()
    document.add_heading(f'Lista - {evento.descricao}', 0)
    document.add_paragraph(f'Gerado em: {now().strftime("%d/%m/%Y %H:%M")}')
    document.add_paragraph('')

    # Congressistas
    document.add_heading('Congressistas', level=1)
    document.add_paragraph(f'Total de congressistas: {total_congressistas}')
    for c in lista_congressistas:
        document.add_paragraph(
            f'{c["nome"]} — {c["cpf"]}',
            style='List Number'
        )

    document.add_paragraph('')

    # Profissionais
    document.add_heading('Profissionais', level=1)
    document.add_paragraph(f'Total de profissionais: {total_profissionais}')
    for p in lista_profissionais:
        document.add_paragraph(
            f'{p["nome"]} — {p["cpf"]}',
            style='List Number'
        )

    document.add_paragraph('')
    document.add_paragraph(f'Total geral (congressistas + profissionais): {total_geral}')

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="evento_{evento.id}_inscritos.docx"'
    document.save(response)
    return response

from django.http import HttpResponse
from django.utils.timezone import now
from docx import Document
from .models import Profissional

class ProfissionalRelatorioDocxView(View):
    """ Gera relatório minimalista apenas com dados essenciais """
    
    def get(self, request, *args, **kwargs):
        profissionais = Profissional.objects.all().order_by('nome')
        
        document = Document()
        document.add_heading('Profissionais', 0)
        document.add_paragraph(now().strftime("%d/%m/%Y %H:%M"))
        document.add_paragraph('')

        for profissional in profissionais:
            eventos = ", ".join([e.descricao for e in profissional.eventos.all()])
            
            document.add_paragraph(
                f'{profissional.nome} • {profissional.cpf or "---"} • {eventos or "---"}'
            )

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="profissionais.docx"'
        document.save(response)
        return response
    
class PlanejamentoRelatorioDocxView(View):
    """ Gera relatório de planejamento com totais e formatação """
    
    def get(self, request, *args, **kwargs):
        planejamentos = Planejamento.objects.all().order_by('descricao')
        total_geral = sum(p.valor_planejado for p in planejamentos if p.valor_planejado)
        
        document = Document()
        document.add_heading('Relatório de Planejamento Financeiro', 0)
        document.add_paragraph(f'Gerado em: {now().strftime("%d/%m/%Y %H:%M")}')
        document.add_paragraph(f'Total de itens: {planejamentos.count()}')
        document.add_paragraph(f'Valor total planejado: R$ {total_geral:,.2f}')
        document.add_paragraph('')

        # Tabela principal
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Cabeçalho
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DESCRIÇÃO'
        hdr_cells[1].text = 'VALOR PLANEJADO'
        hdr_cells[2].text = 'STATUS'
        
        # Dados
        for planejamento in planejamentos:
            row_cells = table.add_row().cells
            row_cells[0].text = planejamento.descricao or "---"
            
            # Formata valor
            if planejamento.valor_planejado:
                row_cells[1].text = f"R$ {planejamento.valor_planejado:,.2f}"
            else:
                row_cells[1].text = "R$ 0,00"
            
            # Formata status
            status = planejamento.status
            if hasattr(planejamento, 'get_status_display'):
                status = planejamento.get_status_display()
            row_cells[2].text = str(status)

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="planejamento_financeiro.docx"'
        document.save(response)
        return response

class PedidosSimplesRelatorioDocxView(View):
    """ Gera relatório minimalista ordenado por nome com resumo """
    
    def get(self, request, *args, **kwargs):
        from django.db.models.functions import Lower
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Ordenar por nome completo em ordem alfabética
        pedidos = PedidoCamisa.objects.all().order_by(Lower('nome_completo'))
        
        # Preparar dados para o resumo
        resumo_por_tipo = {}  # Dicionário para agrupar por tipo de camisa
        
        for pedido in pedidos:
            # Usar o tipo da camisa (choices TIPO_CAMISA) em vez da descrição
            tipo_camisa = pedido.camisa.get_tipo_display() if pedido.camisa else "Sem tipo"
            tamanho = pedido.get_tamanho_display() if hasattr(pedido, 'get_tamanho_display') else pedido.tamanho
            cor = pedido.get_cor_display() if hasattr(pedido, 'get_cor_display') else pedido.cor
            
            # Criar chave combinando tamanho e cor
            chave = f"{tamanho} {cor}"
            
            # Inicializar o dicionário para este tipo de camisa se não existir
            if tipo_camisa not in resumo_por_tipo:
                resumo_por_tipo[tipo_camisa] = {}
            
            # Adicionar ao resumo
            if chave in resumo_por_tipo[tipo_camisa]:
                resumo_por_tipo[tipo_camisa][chave] += 1
            else:
                resumo_por_tipo[tipo_camisa][chave] = 1
        
        document = Document()
        document.add_heading('Pedidos por Ordem Alfabética', 0)
        
        # Data e hora
        data_paragraph = document.add_paragraph()
        data_run = data_paragraph.add_run(now().strftime("%d/%m/%Y %H:%M"))
        data_run.bold = True
        data_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        document.add_paragraph('')
        
        # Adicionar lista de pedidos
        document.add_heading('Lista de Pedidos', level=1)
        
        # Adicionar cabeçalho à lista
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nome'
        hdr_cells[1].text = 'Cidade'
        hdr_cells[2].text = 'Tipo Camisa'
        hdr_cells[3].text = 'Tamanho'
        hdr_cells[4].text = 'Cor'
        hdr_cells[5].text = 'Status'
        
        for pedido in pedidos:
            tamanho = pedido.get_tamanho_display() if hasattr(pedido, 'get_tamanho_display') else pedido.tamanho
            cor = pedido.get_cor_display() if hasattr(pedido, 'get_cor_display') else pedido.cor
            status = pedido.get_status_display() if hasattr(pedido, 'get_status_display') else pedido.status
            tipo_camisa = pedido.camisa.get_tipo_display() if pedido.camisa else "Sem tipo"
            descricao_camisa = pedido.camisa.descricao if pedido.camisa else "Sem descrição"
            
            # Adicionar linha à tabela
            row_cells = table.add_row().cells
            row_cells[0].text = pedido.nome_completo or ''
            row_cells[1].text = pedido.cidade or ''
            row_cells[2].text = f"{tipo_camisa} ({descricao_camisa})"
            row_cells[3].text = tamanho
            row_cells[4].text = cor
            row_cells[5].text = status
        
        document.add_paragraph('')
        
        # Adicionar resumo
        document.add_heading('Resumo por Tipo de Camisa, Tamanho e Cor', level=1)
        
        # DEBUG: Mostrar todos os tipos de camisa encontrados
        tipos = set()
        for pedido in pedidos:
            if pedido.camisa:
                tipos.add(pedido.camisa.get_tipo_display())
        
        debug_p = document.add_paragraph()
        debug_p.add_run('Tipos encontrados: ').bold = True
        debug_p.add_run(', '.join(sorted(tipos)))
        document.add_paragraph('')
        
        # Resumo por tipo de camisa (ordenado alfabeticamente)
        for tipo_camisa in sorted(resumo_por_tipo.keys()):
            resumo = resumo_por_tipo[tipo_camisa]
            document.add_heading(f'Camisas {tipo_camisa}', level=2)
            
            if resumo:
                # Ordenar por tamanho e depois por cor
                itens_ordenados = sorted(resumo.items(), key=lambda x: (x[0].split()[0], x[0].split()[1]))
                
                for chave, quantidade in itens_ordenados:
                    partes = chave.split()
                    tamanho = partes[0]
                    cor = ' '.join(partes[1:])
                    
                    p = document.add_paragraph()
                    p.add_run(f'{quantidade} ').bold = True
                    p.add_run(f'{tamanho} {cor}')
            else:
                document.add_paragraph('Nenhum pedido encontrado')
            
            document.add_paragraph('')
        
        # Adicionar totais
        document.add_heading('Totais Gerais', level=2)
        
        total_por_tipo = {}
        for tipo_camisa, resumo in resumo_por_tipo.items():
            total_por_tipo[tipo_camisa] = sum(resumo.values())
        
        total_geral = sum(total_por_tipo.values())
        
        for tipo_camisa, total in sorted(total_por_tipo.items()):
            p_total = document.add_paragraph()
            p_total.add_run(f'Total {tipo_camisa}: ').bold = True
            p_total.add_run(f'{total}')
        
        p_total = document.add_paragraph()
        p_total.add_run('Total Geral: ').bold = True
        p_total.add_run(f'{total_geral}').bold = True

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="pedidos_com_resumo.docx"'
        document.save(response)
        return response
class CaixaCompletoRelatorioDocxView(View):
    """ Gera relatório completo do caixa sem tabelas """
    
    def get(self, request, *args, **kwargs):
        document = Document()
        document.add_heading('Relatório Completo do Caixa', 0)
        document.add_paragraph(f'Gerado em: {now().strftime("%d/%m/%Y %H:%M")}')
        document.add_paragraph('')
        
        # ============ CÁLCULOS ============
        
        # PLANEJAMENTO
        total_planejado = Planejamento.objects.aggregate(
            total=Sum('valor_planejado')
        )['total'] or 0
        
        # PAGAMENTOS (CAIXA GERAL)
        total_entradas = Pagamento.objects.filter(
            valor_pago__gt=0
        ).aggregate(total=Sum('valor_pago'))['total'] or 0
        
        total_saidas = Pagamento.objects.filter(
            valor_pago__lt=0
        ).aggregate(total=Sum('valor_pago'))['total'] or 0
        
        total_saidas = abs(total_saidas) if total_saidas else 0
        saldo_caixa = total_entradas - total_saidas
        
        # INSCRIÇÕES
        total_inscricoes = Inscricao.objects.aggregate(
            total=Sum('valor_total')
        )['total'] or 0
        
        inscricao_content_type = ContentType.objects.get_for_model(Inscricao)
        pagamentos_inscricoes = Pagamento.objects.filter(
            content_type=inscricao_content_type,
            valor_pago__gt=0
        ).aggregate(total=Sum('valor_pago'))['total'] or 0
        
        total_inscricoes_pagas = pagamentos_inscricoes
        total_inscricoes_receber = max(total_inscricoes - pagamentos_inscricoes, 0)
        
        # CAMISAS
        total_camisas = PedidoCamisa.objects.aggregate(
            total=Sum('valor_venda')
        )['total'] or 0
        
        camisas_pagas = PedidoCamisa.objects.filter(
            status='pago'
        ).aggregate(total=Sum('valor_venda'))['total'] or 0
        
        camisas_receber = total_camisas - camisas_pagas
        
        # SALDO FUTURO PREVISTO
        saldo_futuro_previsto = saldo_caixa + total_inscricoes_receber + camisas_receber - total_planejado
        
        # ============ RELATÓRIO ============
        
        # RESUMO GERAL
        document.add_heading('Resumo Financeiro Geral', level=1)
        
        # CAIXA ATUAL
        p = document.add_paragraph()
        p.add_run('Valor no Caixa Atual: ').bold = True
        p.add_run(f'R$ {saldo_caixa:,.2f}')
        
        document.add_paragraph('')
        
        # PLANEJAMENTO
        document.add_heading('Planejamento', level=2)
        p = document.add_paragraph()
        p.add_run('Total Planejado: ').bold = True
        p.add_run(f'R$ {total_planejado:,.2f}')
        
        document.add_paragraph('')
        
        # FLUXO DE CAIXA
        document.add_heading('Fluxo de Caixa', level=2)
        p = document.add_paragraph()
        p.add_run('Total de Entradas: ').bold = True
        p.add_run(f'R$ {total_entradas:,.2f}')
        
        p = document.add_paragraph()
        p.add_run('Total de Saídas: ').bold = True
        p.add_run(f'R$ {total_saidas:,.2f}')
        
        document.add_paragraph('')
        
        # INSCRIÇÕES
        document.add_heading('Inscrições', level=2)
        p = document.add_paragraph()
        p.add_run('Valor Total das Inscrições: ').bold = True
        p.add_run(f'R$ {total_inscricoes:,.2f}')
        
        p = document.add_paragraph()
        p.add_run('Inscrições Pagas: ').bold = True
        p.add_run(f'R$ {total_inscricoes_pagas:,.2f}')
        
        p = document.add_paragraph()
        p.add_run('Total a Receber (Inscrições): ').bold = True
        p.add_run(f'R$ {total_inscricoes_receber:,.2f}')
        
        document.add_paragraph('')
        
        # CAMISAS
        document.add_heading('Camisas', level=2)
        p = document.add_paragraph()
        p.add_run('Vendas de Camisas: ').bold = True
        p.add_run(f'R$ {total_camisas:,.2f}')
        
        p = document.add_paragraph()
        p.add_run('Camisas Pagas: ').bold = True
        p.add_run(f'R$ {camisas_pagas:,.2f}')
        
        p = document.add_paragraph()
        p.add_run('Total a Receber (Camisas): ').bold = True
        p.add_run(f'R$ {camisas_receber:,.2f}')
        
        document.add_paragraph('')
        
        # PROJEÇÃO FUTURA
        document.add_heading('Projeção Futura', level=2)
        p = document.add_paragraph()
        p.add_run('Saldo Futuro Previsto: ').bold = True
        p.add_run(f'R$ {saldo_futuro_previsto:,.2f}')
        
        document.add_paragraph('')
        document.add_paragraph('')
        
        # ÚLTIMOS MOVIMENTOS
        document.add_heading('Últimos Movimentos do Caixa', level=1)
        
        ultimos_movimentos = Pagamento.objects.order_by('-data_pagamento')[:10]
        
        if ultimos_movimentos:
            for movimento in ultimos_movimentos:
                p = document.add_paragraph()
                
                # Data
                data_str = movimento.data_pagamento.strftime("%d/%m/%Y") if movimento.data_pagamento else "---"
                
                # Tipo
                tipo = "ENTRADA" if movimento.valor_pago > 0 else "SAÍDA"
                
                # Origem
                origem = "Geral"
                if movimento.content_type:
                    origem = f"{movimento.content_type.model.capitalize()}"
                    if movimento.object_id:
                        origem += f" #{movimento.object_id}"
                
                # Valor
                valor = abs(movimento.valor_pago)
                
                p.add_run(f'{data_str} - ').bold = True
                p.add_run(f'{tipo} - ')
                p.add_run(f'{origem} - ')
                p.add_run(f'R$ {valor:,.2f}')
        else:
            document.add_paragraph('Nenhum movimento encontrado.')

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="relatorio_caixa_simples.docx"'
        document.save(response)
        return response
from rest_framework.response import Response

# View para processar pagamento de inscrição via gateway
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .services.payment_service import PaymentService

@api_view(['POST'])
def processar_pagamento_inscricao(request):
    """
    Endpoint para processar pagamento de inscrição
    Suporta cartão de crédito e PIX
    """
    inscricao_id = request.data.get('inscricao_id')
    gateway = request.data.get('gateway', 'pagseguro')
    tipo_pagamento = request.data.get('tipo_pagamento', 'credit_card')
    dados_pagamento = request.data.get('dados_pagamento', {})
    
    if not inscricao_id:
        return Response(
            {'error': 'inscricao_id é obrigatório'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    try:
        inscricao = Inscricao.objects.get(id=inscricao_id)
    except Inscricao.DoesNotExist:
        return Response(
            {'error': 'Inscrição não encontrada'},
            status=status.HTTP_404_NOT_FOUND
        )
    
    resultado = PaymentService.pagar_inscricao(inscricao, tipo_pagamento, dados_pagamento, gateway)
    
    if resultado['success']:
        return Response(resultado, status=status.HTTP_201_CREATED)
    else:
        return Response(resultado, status=status.HTTP_400_BAD_REQUEST)
